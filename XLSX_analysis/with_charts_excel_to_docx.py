# -*- coding: utf-8 -*-
"""
excel_to_single.py

将 Excel 的“表格纯文本（含合并单元格展开）+ 图表纯文本摘要（Line/Bar/Pie）”
合并输出到**单个文件**（自动按扩展名选择 .docx / .md / .txt）。

用法：
# 
# 生成单个 Word 文档（表格纯文本 + 图表摘要 全都合在一起）python with_charts_excel_to_docx.py with_charts_exal_data/data.xlsx --out with_charts_exal_solution_result/withchars.docx
  # 只要表格，不要图表 python excel_to_docx.py 扁平化完整数据结果.xlsx --out result.docx --no-charts
"""

import argparse
import math
from pathlib import Path
from typing import List, Tuple, Any, Optional, Dict

import pandas as pd
from openpyxl import load_workbook
from openpyxl.utils.cell import range_boundaries, coordinate_from_string, column_index_from_string

# docx 仅在输出为 .docx 时才需要
try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


# --------------------------- 工具 ---------------------------

def _format_number(v: Any) -> str:
    if isinstance(v, (int, float)) and v is not None:
        try:
            if isinstance(v, float) and (math.isnan(v) or math.isinf(v)):
                return str(v)
        except Exception:
            pass
        return f"{v:.6g}"
    return str(v)


def _shorten_points(points: List[Tuple[Any, Any]], max_points: int = 20) -> str:
    if len(points) <= max_points:
        return "；".join([f"{x} → {y}" for x, y in points])
    head = "；".join([f"{x} → {y}" for x, y in points[:10]])
    tail = "；".join([f"{x} → {y}" for x, y in points[-3:]])
    return f"{head}；…（共 {len(points)} 点）…；{tail}"


def _read_ref_string_values(wb, ref_str: str):
    """
    从 'Sheet'!$A$1:$A$5 读取值列表及边界。
    返回 (sheet_name, values, (min_col, min_row, max_col, max_row))
    """
    sheet_name, rng = ref_str.split("!", 1)
    sheet_name = sheet_name.strip().strip("'").strip('"')
    ws = wb[sheet_name]
    rng = rng.replace("$", "")

    if ":" in rng:
        min_cell, max_cell = rng.split(":")
    else:
        min_cell = max_cell = rng

    col_letters_min, row_min = coordinate_from_string(min_cell)
    col_letters_max, row_max = coordinate_from_string(max_cell)
    min_col = column_index_from_string(col_letters_min)
    max_col = column_index_from_string(col_letters_max)
    min_row = row_min
    max_row = row_max

    vals = []
    for r in range(min_row, max_row + 1):
        row_vals = []
        for c in range(min_col, max_col + 1):
            v = ws.cell(row=r, column=c).value
            row_vals.append(v)
        if min_col == max_col or min_row == max_row:
            vals.extend(row_vals)
        else:
            vals.append(row_vals)

    return sheet_name, vals, (min_col, min_row, max_col, max_row)


def _read_range_values(wb, ref) -> Tuple[str, List[Any]]:
    try:
        ws = wb[ref.sheetname] if hasattr(ref, "sheetname") else wb[ref.sheet]
    except KeyError:
        ws = wb.active
    min_col, min_row, max_col, max_row = range_boundaries(ref.range)
    vals = []
    for r in range(min_row, max_row + 1):
        row_vals = []
        for c in range(min_col, max_col + 1):
            row_vals.append(ws.cell(row=r, column=c).value)
        if min_col == max_col or min_row == max_row:
            vals.extend(row_vals)
        else:
            vals.append(row_vals)
    return (ws.title, vals)


def _safe_title(obj: Any) -> Optional[str]:
    if obj is None:
        return None
    if isinstance(obj, str):
        s = obj.strip()
        return s or None
    try:
        s = str(obj).strip()
        return s or None
    except Exception:
        return None


def _extract_title_from_tx(wb, tx_obj) -> Optional[str]:
    if tx_obj is None:
        return None
    strRef = getattr(tx_obj, "strRef", None)
    if strRef and getattr(strRef, "f", None):
        try:
            _, vals, _ = _read_ref_string_values(wb, strRef.f)
            return str(vals[0]) if vals else None
        except Exception:
            pass
    rich = getattr(tx_obj, "rich", None)
    if rich and getattr(rich, "paragraphs", None):
        texts = []
        for para in rich.paragraphs:
            runs = getattr(para, "r", None) or getattr(para, "runs", None) or []
            for run in runs:
                t = getattr(run, "t", None)
                if t:
                    texts.append(t)
        if texts:
            s = "".join(texts).strip()
            if s:
                return s
    return _safe_title(tx_obj)


def _chart_title(wb, ch) -> Optional[str]:
    t = getattr(ch, "title", None)
    if t is None:
        return None
    tx = getattr(t, "tx", None)
    if tx:
        return _extract_title_from_tx(wb, tx)
    return _safe_title(t)


def _series_title(wb, s, val_ref_str: Optional[str]) -> Optional[str]:
    tx = getattr(s, "tx", None)
    if tx:
        title = _extract_title_from_tx(wb, tx)
        if title:
            return title
    if val_ref_str:
        try:
            sheet, _, (min_c, min_r, _, _) = _read_ref_string_values(wb, val_ref_str)
            hdr_row = min_r - 1
            if hdr_row >= 1:
                ws = wb[sheet]
                hdr = ws.cell(row=hdr_row, column=min_c).value
                if hdr is not None:
                    return str(hdr)
        except Exception:
            pass
    return None


# --------------------------- 图表解析 ---------------------------

def extract_charts_as_text_by_sheet(wb) -> Dict[str, List[str]]:
    out: Dict[str, List[str]] = {}
    for ws in wb.worksheets:
        charts = getattr(ws, "_charts", [])
        if not charts:
            continue

        blocks: List[str] = []
        for idx, ch in enumerate(charts, 1):
            chart_type = type(ch).__name__
            if chart_type == "LineChart":
                ctype_cn = "折线图"
            elif chart_type == "BarChart":
                orient = getattr(ch, "type", None)
                ctype_cn = "柱状图" if orient == "col" else "条形图"
            elif chart_type == "PieChart":
                ctype_cn = "饼图"
            else:
                ctype_cn = f"{chart_type}（其他类型）"

            chart_title = _chart_title(wb, ch) or "（无标题）"

            categories_values = None
            if hasattr(ch, "categories") and ch.categories is not None:
                try:
                    _, categories_values = _read_range_values(wb, ch.categories)
                except Exception:
                    categories_values = None

            series_texts: List[str] = []

            for s_i, s in enumerate(getattr(ch, "series", []), 1):
                val_ref = None
                if getattr(s, "val", None) and getattr(s.val, "numRef", None) and getattr(s.val.numRef, "ref", None):
                    val_ref = s.val.numRef.ref

                cat_vals = None
                if getattr(s, "cat", None):
                    if getattr(s.cat, "strRef", None) and getattr(s.cat.strRef, "ref", None):
                        try:
                            _, cat_vals, _ = _read_ref_string_values(wb, s.cat.strRef.ref)
                        except Exception:
                            cat_vals = None
                    elif getattr(s.cat, "numRef", None) and getattr(s.cat.numRef, "ref", None):
                        try:
                            _, cat_vals, _ = _read_ref_string_values(wb, s.cat.numRef.ref)
                        except Exception:
                            cat_vals = None

                yvals = None
                if val_ref:
                    try:
                        _, yvals, _ = _read_ref_string_values(wb, val_ref)
                    except Exception:
                        yvals = None

                s_title = _series_title(wb, s, val_ref) or f"序列{s_i}"

                points = []
                if yvals is not None:
                    xs = cat_vals if cat_vals is not None else categories_values
                    if xs is not None:
                        for x, y in zip(xs, yvals):
                            points.append((x, y))
                    else:
                        for i, y in enumerate(yvals, 1):
                            points.append((i, y))

                if chart_type == "PieChart":
                    if points:
                        nums = [p[1] for p in points if isinstance(p[1], (int, float)) and p[1] is not None]
                        total = sum(nums) if nums else None
                        items = []
                        for x, y in points:
                            if total and isinstance(y, (int, float)):
                                pct = f"{y * 100 / total:.2f}%"
                            else:
                                pct = "—"
                            items.append(f"{x}：{_format_number(y)}（{pct}）")
                        series_texts.append(f"【{s_title}】" + "；".join(items))
                    else:
                        series_texts.append(f"【{s_title}】无有效数据")
                else:
                    if points:
                        series_texts.append(
                            f"【{s_title}】" + _shorten_points([(x, _format_number(y)) for x, y in points])
                        )
                    else:
                        series_texts.append(f"【{s_title}】无有效数据")

            lines = [
                f"第 {idx} 个图表：{ctype_cn}",
                f"标题：{chart_title}",
            ]
            if series_texts:
                lines.append("数据：")
                lines.extend(["  " + s for s in series_texts])
            else:
                lines.append("数据：无")

            blocks.append("\n".join(lines))

        if blocks:
            out[ws.title] = blocks

    return out


# --------------------------- 表格展开 ---------------------------

def unmerge_excel_cells_from_wb(wb) -> Dict[str, pd.DataFrame]:
    sheet_data: Dict[str, pd.DataFrame] = {}
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        max_row, max_col = ws.max_row, ws.max_column

        matrix = [[None for _ in range(max_col)] for _ in range(max_row)]
        for r in range(1, max_row + 1):
            for c in range(1, max_col + 1):
                matrix[r - 1][c - 1] = ws.cell(r, c).value

        for merged_range in ws.merged_cells.ranges:
            min_row, min_col, max_row_r, max_col_r = (
                merged_range.min_row,
                merged_range.min_col,
                merged_range.max_row,
                merged_range.max_col,
            )
            value = ws.cell(min_row, min_col).value
            for r in range(min_row - 1, max_row_r):
                for c in range(min_col - 1, max_col_r):
                    matrix[r][c] = value

        df = pd.DataFrame(matrix)
        df.dropna(axis=0, how='all', inplace=True)
        df.dropna(axis=1, how='all', inplace=True)
        df.fillna("", inplace=True)
        sheet_data[sheet] = df

    return sheet_data


# --------------------------- 单文件输出 ---------------------------

def write_single_output(
    out_path: str,
    sheet_data: Dict[str, pd.DataFrame],
    charts_by_sheet: Dict[str, List[str]],
):
    ext = Path(out_path).suffix.lower()

    if ext == ".docx":
        if not DOCX_AVAILABLE:
            raise RuntimeError("未安装 python-docx，无法写入 .docx。请先 pip install python-docx。")
        doc = Document()

        for sheet_name, df in sheet_data.items():
            # Sheet 标题
            doc.add_heading(f"【Sheet】{sheet_name}", level=1)

            # 计算列宽用于等宽文本（写段落；如需真正 Word 表格可改用 add_table）
            col_widths = [max(len(str(v)) for v in df[col]) for col in df.columns] if df.shape[1] else []

            # 表格纯文本
            for _, row in df.iterrows():
                non_empty = [str(cell).strip() for cell in row if str(cell).strip()]
                if not non_empty:
                    continue
                if len(set(non_empty)) == 1:
                    line = non_empty[0]
                else:
                    formatted = [str(cell).ljust(col_widths[i]) for i, cell in enumerate(row)]
                    line = "  ".join(formatted).rstrip()
                doc.add_paragraph(line)

            # 图表纯文本（直接接在表格后面）
            if sheet_name in charts_by_sheet and charts_by_sheet[sheet_name]:
                doc.add_paragraph("")  # 空行
                doc.add_heading("图表文字摘要", level=2)
                for block in charts_by_sheet[sheet_name]:
                    doc.add_paragraph(block)
                    doc.add_paragraph("")

        doc.save(out_path)
        print(f"✅ 已生成：{out_path}")

    else:
        # .md 或 .txt：按行写文本
        lines: List[str] = []
        for sheet_name, df in sheet_data.items():
            lines.append(f"# 【Sheet】{sheet_name}")
            lines.append("")

            col_widths = [max(len(str(v)) for v in df[col]) for col in df.columns] if df.shape[1] else []

            for _, row in df.iterrows():
                non_empty = [str(cell).strip() for cell in row if str(cell).strip()]
                if not non_empty:
                    continue
                if len(set(non_empty)) == 1:
                    line = non_empty[0]
                else:
                    formatted = [str(cell).ljust(col_widths[i]) for i, cell in enumerate(row)]
                    line = "  ".join(formatted).rstrip()
                lines.append(line)

            lines.append("")
            # 图表纯文本（紧接着输出）
            if sheet_name in charts_by_sheet and charts_by_sheet[sheet_name]:
                lines.append("## 图表文字摘要")
                lines.append("")
                for block in charts_by_sheet[sheet_name]:
                    lines.extend(block.splitlines())
                    lines.append("")
            lines.append("")

        Path(out_path).write_text("\n".join(lines), encoding="utf-8")
        print(f"✅ 已生成：{out_path}")


# --------------------------- 主流程 ---------------------------

def main():
    parser = argparse.ArgumentParser(description="将 Excel 的表格纯文本与图表纯文本合并到单个输出文件")
    parser.add_argument("excel", help="输入 Excel 文件（.xlsx）")
    parser.add_argument("--out", required=True, help="唯一输出文件路径（.docx / .md / .txt）")
    parser.add_argument("--no-charts", action="store_true", help="不解析图表（只输出表格纯文本）")
    args = parser.parse_args()

    xlsx = Path(args.excel)
    if not xlsx.exists():
        raise FileNotFoundError(f"Excel 文件不存在：{xlsx}")

    wb = load_workbook(xlsx, data_only=True, read_only=False)

    # 1) 表格
    sheet_data = unmerge_excel_cells_from_wb(wb)

    # 2) 图表
    charts_by_sheet = {} if args.no_charts else extract_charts_as_text_by_sheet(wb)

    # 3) 单文件写出
    write_single_output(args.out, sheet_data, charts_by_sheet)

    wb.close()


if __name__ == "__main__":
    main()