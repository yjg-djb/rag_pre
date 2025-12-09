import warnings
warnings.filterwarnings('ignore', category=FutureWarning)
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import load_workbook
from pathlib import Path
from datetime import datetime
import re
import numpy as np
import os
import logging

# -------------------------- 日志配置（日志文件保存到logs文件夹） --------------------------
def setup_logger(log_file="logs/excel转Word批量日志.log"):
    """配置日志：同时输出到控制台和logs文件夹下的日志文件"""
    logger = logging.getLogger("ExcelToDocxBatch")
    logger.setLevel(logging.INFO)
    
    if logger.handlers:
        return logger
    
    # 确保logs文件夹存在
    log_dir = os.path.dirname(log_file)
    if log_dir and not os.path.exists(log_dir):
        os.makedirs(log_dir, exist_ok=True)
        logger.info(f"已自动创建日志文件夹：{os.path.abspath(log_dir)}")
    
    # 日志格式：时间-级别-消息
    formatter = logging.Formatter(
        '%(asctime)s - %(levelname)s - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    
    # 控制台处理器
    console_handler = logging.StreamHandler()
    console_handler.setFormatter(formatter)
    logger.addHandler(console_handler)
    
    # 文件处理器（支持中文，保存到logs文件夹）
    file_handler = logging.FileHandler(log_file, encoding='utf-8')
    file_handler.setFormatter(formatter)
    logger.addHandler(file_handler)
    
    return logger

logger = setup_logger()

# -------------------------- 核心工具函数 --------------------------
def parse_cell_date(cell_value, cell_number_format, date_formats=None):
    """精准解析日期格式"""
    if date_formats is None:
        date_formats = ["%Y-%m-%d %H:%M:%S", "%Y-%m-%d"]
    
    if isinstance(cell_value, datetime):
        if cell_value.tzinfo:
            return cell_value.strftime(f"{date_formats[0]}%z")
        return cell_value.strftime(date_formats[0])
    
    date_patterns = r'[yYmMdDhHms]'
    if isinstance(cell_value, (int, float)) and re.search(date_patterns, cell_number_format):
        try:
            base_date = datetime(1899, 12, 30) if cell_value > 60 else datetime(1900, 1, 1)
            date_val = base_date + pd.Timedelta(days=cell_value)
            return date_val.strftime(date_formats[0])
        except:
            pass
    
    for fmt in date_formats:
        try:
            return datetime.strptime(str(cell_value), fmt).strftime(fmt)
        except:
            continue
    
    return str(cell_value) if cell_value is not None else ""

def unmerge_and_fill_excel(xlsx_path, output_path=None):
    """处理Excel：取消所有合并单元格并填充值"""
    xlsx_file = Path(xlsx_path)
    if not output_path:
        output_path = xlsx_file.parent / f"{xlsx_file.stem}_处理后.xlsx"
    
    # 加载工作簿
    wb = load_workbook(xlsx_path)
    
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        # 复制合并区域信息（避免迭代时修改集合导致错误）
        merged_ranges = list(ws.merged_cells.ranges)
        
        for merged_range in merged_ranges:
            # 获取合并区域的边界
            min_row, min_col, max_row, max_col = (
                merged_range.min_row,
                merged_range.min_col,
                merged_range.max_row,
                merged_range.max_col
            )
            
            # 获取左上角单元格的值和格式
            top_left_cell = ws.cell(row=min_row, column=min_col)
            fill_value = top_left_cell.value
            fill_number_format = top_left_cell.number_format
            
            # 取消合并
            ws.unmerge_cells(str(merged_range))
            
            # 填充所有单元格
            for row in range(min_row, max_row + 1):
                for col in range(min_col, max_col + 1):
                    cell = ws.cell(row=row, column=col)
                    cell.value = fill_value
                    cell.number_format = fill_number_format
    
    # 保存处理后的Excel
    wb.save(output_path)
    wb.close()
    return output_path

def format_table_as_text(data, headers):
    """将表格数据格式化为按产品分组的纯文本形式（每行一个产品）"""
    result = []
    result.append("[表格内容]:")
    result.append("")
    
    # 直接处理数据（不需要跳过标题行，因为已经在xlsx_to_docx中处理了）
    for row_idx, row in enumerate(data, 1):
        product_info = []
        
        # 跳过"原始行号"列，从第二列开始处理
        for col_idx, (header, cell_value) in enumerate(zip(headers[1:], row[1:])):
            cell_str = str(cell_value).strip() if cell_value else ""
            
            if cell_str:  # 只处理非空单元格
                product_info.append(f"{header}: {cell_str}")
        
        if product_info:
            # 用中文逗号连接所有字段，形成一个段落
            result.append("，".join(product_info))
    
    return "\n".join(result)

def xlsx_to_docx(
    xlsx_path,
    docx_save_path=None,
    processed_xlsx_path=None,
    keep_merge_info=False,
    formula_mode="result",
    date_formats=None,
    filter_empty_rows=True,
    split_by_row=False,
    chunk_size=500,
    font_size=10,
    output_format="table"  # 输出格式 "table" 或 "text"
):
    """基于处理后的Excel生成DOCX"""
    xlsx_file = Path(xlsx_path)
    if docx_save_path:
        base_save_path = Path(docx_save_path)
    else:
        base_save_path = xlsx_file.parent / f"{xlsx_file.stem}_纯文本.docx"

    # 使用处理后的Excel进行转换
    use_xlsx_path = processed_xlsx_path if processed_xlsx_path else xlsx_path
    wb_result = load_workbook(use_xlsx_path, data_only=True)
    wb_formula = load_workbook(use_xlsx_path, data_only=False) if formula_mode != "result" else None

    # 创建文档
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(font_size)

    for sheet_idx, sheet_name in enumerate(wb_result.sheetnames, 1):
        ws_result = wb_result[sheet_name]
        ws_formula = wb_formula[sheet_name] if wb_formula and sheet_name in wb_formula.sheetnames else None

        data = []
        original_row_nums = []
        max_col_count = 0
        
        for row_idx, row in enumerate(ws_result.iter_rows(), 1):
            current_row = []
            for col_idx, cell in enumerate(row, 1):
                cell_result = cell.value
                cell_formula = ws_formula.cell(row=row_idx, column=col_idx).value if ws_formula else None
                cell_format = cell.number_format

                if formula_mode == "result":
                    val = cell_result
                elif formula_mode == "formula":
                    val = cell_formula
                else:
                    val = f"[{cell_formula}]={cell_result}" if cell_formula else cell_result

                parsed_val = parse_cell_date(val, cell_format, date_formats)
                current_row.append(parsed_val if parsed_val else "")
            
            if len(current_row) > max_col_count:
                max_col_count = len(current_row)
            
            # 检查是否为空行
            if filter_empty_rows and all(not str(cell).strip() for cell in current_row):
                continue
            data.append(current_row)
            original_row_nums.append(row_idx)

        # 补齐列数
        for row in data:
            while len(row) < max_col_count:
                row.append("")

        # 使用第一行数据作为表头（如果存在）
        if len(data) > 0:
            headers = ["原始行号"] + [str(cell).strip() if cell else f"列{i+1}" for i, cell in enumerate(data[0])]
            # 移除第一行（表头行）从数据中
            data = data[1:]
            original_row_nums = original_row_nums[1:]
        else:
            headers = ["原始行号"] + [f"列{col_idx+1}" for col_idx in range(max_col_count)]

        # 添加Sheet标题
        if sheet_idx > 1:
            doc.add_page_break()
        
        title = doc.add_heading(f'{sheet_idx}、{sheet_name}', level=1)
        
        if output_format == "text":
            # 纯文本格式输出（按产品分组，每行一个产品）
            formatted_text = format_table_as_text(
                [[str(row_num)] + row for row_num, row in zip(original_row_nums, data)],
                headers
            )
            doc.add_paragraph(formatted_text)
        else:
            # 原有的表格格式输出
            info_para = doc.add_paragraph()
            info_para.add_run(f"数据维度：{len(data)}行 × {max_col_count}列\n").bold = True
            doc.add_paragraph("-" * 80)

            header_text = "\t".join(headers)
            header_para = doc.add_paragraph(header_text)
            header_para.bold = True
            doc.add_paragraph("-" * 80)

            for original_row, data_row in zip(original_row_nums, data):
                row_text = f"{original_row}\t" + "\t".join(data_row)
                doc.add_paragraph(row_text)

    # 保存文档
    doc.save(str(base_save_path))
    logger.info(f"文档保存完成：{base_save_path}")

    wb_result.close()
    if wb_formula:
        wb_formula.close()

    logger.info("当前Excel所有Sheet处理完成！")

if __name__ == "__main__":
    # -------------------------- 批量处理配置 --------------------------
    excel_dir = "exal_solution_result"  # 待处理Excel文件存放目录
    output_docx_dir = "docx_result"     # Word输出目录
    log_dir = "logs"                    # 日志目录（自动创建）
    
    # -------------------------- 执行批量处理 --------------------------
    logger.info("="*80)
    logger.info("🚀 启动Excel转Word批量处理程序")
    logger.info(f"📂 待处理Excel文件夹：{os.path.abspath(excel_dir)}")
    logger.info(f"📂 Word输出文件夹：{os.path.abspath(output_docx_dir)}")
    logger.info(f"📜 日志文件夹：{os.path.abspath(log_dir)}")
    logger.info("="*80)
    
    # 确保输出文件夹和日志文件夹存在
    for dir_path in [output_docx_dir, log_dir]:
        if not os.path.exists(dir_path):
            os.makedirs(dir_path, exist_ok=True)
            logger.info(f"✅ 已自动创建文件夹：{os.path.abspath(dir_path)}")
    
    # 检查待处理文件夹是否存在
    if not os.path.exists(excel_dir):
        logger.error(f"❌ 错误：待处理文件夹【{excel_dir}】不存在！")
        logger.error("请确认文件夹名称和路径是否正确，将Excel文件放入该文件夹后重新运行")
        exit(1)
    
    # 筛选Excel文件（仅处理.xlsx格式）
    excel_files = [f for f in os.listdir(excel_dir) if f.endswith('.xlsx')]
    if not excel_files:
        logger.warning(f"⚠️  文件夹【{excel_dir}】中未找到.xlsx格式的Excel文件")
        logger.warning("请检查文件格式是否正确，确保文件未被隐藏")
        exit(0)
    
    # 遍历处理每个Excel文件
    logger.info(f"✅ 找到 {len(excel_files)} 个Excel文件，开始批量处理...")
    for idx, file_name in enumerate(excel_files, 1):
        excel_path = os.path.join(excel_dir, file_name)
        logger.info(f"\n===== 处理进度：{idx}/{len(excel_files)} - 文件：{file_name} =====")
        
        # 第一步：处理Excel（取消合并单元格并填充值）
        processed_xlsx = unmerge_and_fill_excel(excel_path)
        logger.info(f"  处理后的Excel保存路径：{processed_xlsx}")
        
        # 第二步：生成Word文件
        docx_file_name = f"{os.path.splitext(file_name)[0]}_转Word.docx"
        docx_save_path = os.path.join(output_docx_dir, docx_file_name)
        
        try:
            xlsx_to_docx(
                xlsx_path=excel_path,
                processed_xlsx_path=processed_xlsx,
                docx_save_path=docx_save_path,
                keep_merge_info=False,
                formula_mode="result",
                date_formats=["%Y/%m/%d", "%H:%M"],
                filter_empty_rows=True,
                split_by_row=False,  # 改为False，不分块
                chunk_size=100,
                font_size=10,
                output_format="text"  # 使用纯文本格式输出
            )
            logger.info(f"  Word文件生成成功：{docx_save_path}")
        except Exception as e:
            logger.error(f"  处理文件{file_name}时出错：{str(e)}", exc_info=True)
    
    logger.info("\n" + "="*80)
    logger.info("🎉 所有Excel文件批量处理完成！")
    logger.info(f"📊 处理统计：共处理 {len(excel_files)} 个文件")
    logger.info(f"📁 结果文件：保存在【{os.path.abspath(output_docx_dir)}】文件夹")
    logger.info(f"📜 日志文件：保存在【{os.path.abspath(log_dir)}】文件夹")
    logger.info("="*80)