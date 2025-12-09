import os
import re
import shutil
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from paddleocr import PaddleOCR
from PIL import Image, ImageEnhance
import mammoth
from bs4 import BeautifulSoup

# 导入其他模块
from config import INPUT_DIR, OUTPUT_DIR, IMAGE_DIR, OCR_DIR, LOG_DIR
from log import logger
from llm_client import optimize_ocr_with_llm

# 全局初始化
OCR_ENGINE = PaddleOCR(lang='ch')
USED_PREFIXES = set()


# 通用工具函数
def create_required_dirs():
    """创建所有必要目录"""
    for dir_path in [INPUT_DIR, OUTPUT_DIR, IMAGE_DIR, OCR_DIR, LOG_DIR]:
        Path(dir_path).mkdir(parents=True, exist_ok=True)


def get_unique_filepath(base_path: Path) -> Path:
    """生成唯一路径避免冲突"""
    if not base_path.exists():
        return base_path
    counter = 1
    while True:
        new_path = base_path.with_stem(f"{base_path.stem}({counter})")
        if not new_path.exists():
            return new_path
        counter += 1


def sanitize_filename(name: str) -> str:
    """清理文件名非法字符"""
    return re.sub(r'[\\/*?:"<>|]', '_', name)


def get_unique_prefix(base_name: str) -> str:
    """生成唯一前缀"""
    prefix = base_name[:50]
    unique_prefix = prefix
    counter = 1
    while unique_prefix in USED_PREFIXES:
        unique_prefix = f"{prefix}_{counter}"
        counter += 1
    USED_PREFIXES.add(unique_prefix)
    return unique_prefix


# 核心处理函数
def extract_table_text(docx_path: Path) -> list:
    """提取表格文本"""
    try:
        with open(docx_path, "rb") as f:
            html = mammoth.convert_to_html(f).value
        soup = BeautifulSoup(html, 'html.parser')
        table_texts = []

        for table in soup.find_all('table'):
            rows = [row.find_all(['td', 'th']) for row in table.find_all('tr')]
            table_data = [[cell.get_text(strip=True) for cell in row] for row in rows if
                          any(cell.get_text(strip=True) for cell in row)]
            if not table_data:
                continue

            if len(table_data) > 1:
                header, rows_data = table_data[0], table_data[1:]
                table_text = "\n".join([", ".join(
                    [f"{header[i]}: {row[i]}" if i < len(row) else f"{header[i]}: " for i in range(len(header))]) for
                    row in rows_data])
            else:
                table_text = "\n".join([", ".join(row) for row in table_data])
            table_texts.append(table_text)
        return table_texts
    except Exception as e:
        logger.error(f"表格提取失败 {docx_path.name}: {str(e)}")
        return []


def parse_doc_elements(docx_path: Path) -> list:
    """解析文档元素"""
    try:
        table_texts = extract_table_text(docx_path)
        doc = Document(docx_path)
        elements = []
        table_idx = 0

        for elem in doc._element.iter():
            if elem.tag.endswith('}p') and not any(parent.tag.endswith('}tbl') for parent in elem.iterancestors()):
                text = ''.join(run.text for run in elem.iter(qn('w:r')) if run.text).strip()
                if text:
                    elements.append(('text', text))
                for blip in elem.iter(qn('a:blip')):
                    if rId := blip.get(qn('r:embed')):
                        elements.append(('image', rId))
            elif elem.tag.endswith('}tbl') and table_idx < len(table_texts):
                elements.append(('table', table_texts[table_idx]))
                table_idx += 1
        return elements
    except Exception as e:
        logger.error(f"文档解析失败 {docx_path.name}: {str(e)}")
        raise


def extract_images(docx_path: Path, prefix: str) -> dict:
    """提取图片"""
    doc = Document(docx_path)
    image_rels = {rel.rId: rel.target_part.blob for rel in doc.part.rels.values() if "image" in rel.target_ref}
    rid_to_img = {}

    for i, (rId, blob) in enumerate(image_rels.items(), 1):
        base_path = Path(IMAGE_DIR) / f"{prefix}_img_{i:03d}.png"
        img_path = get_unique_filepath(base_path)
        with open(img_path, "wb") as f:
            f.write(blob)
        rid_to_img[rId] = str(img_path)

    logger.info(f"提取图片 {len(rid_to_img)} 张")
    return rid_to_img


def enhance_image(img_path: str) -> str:
    """增强图片对比度"""
    try:
        img = Image.open(img_path)
        enhanced = ImageEnhance.Contrast(img).enhance(2.0)
        enhanced_path = get_unique_filepath(Path(img_path).with_stem(f"{Path(img_path).stem}_enhanced"))
        enhanced.save(enhanced_path)
        logger.info(f"增强图片对比度: {Path(img_path).name}")
        return str(enhanced_path)
    except Exception as e:
        logger.error(f"图片增强失败: {e}")
        return img_path


def ocr_images(image_map: dict, prefix: str) -> tuple[dict, list[float], list[str], int, list[str]]:
    """OCR识别图片文字"""
    rid_to_ocr = {}
    all_confidences = []
    failed_imgs = []
    llm_failed_imgs = []
    success_img_count = 0

    for rId, img_path in image_map.items():
        img_file = Path(img_path)
        img_short_name = img_file.name.split("_")[-1]
        logger.info(f"OCR识别: {img_file.name}")

        if "_img_001.png" in img_path:
            img_path = enhance_image(img_path)

        try:
            result = OCR_ENGINE.predict(img_path)[0]
            texts = []
            if isinstance(result, dict):
                rec_texts = result.get('rec_texts', [])
                rec_scores = result.get('rec_scores', [])
                for t, s in zip(rec_texts, rec_scores):
                    all_confidences.append(s)
                    if s > 0.05:
                        texts.append(str(t))
            else:
                for line in result:
                    if len(line) >= 2:
                        t = line[1][0]
                        s = line[1][1]
                        all_confidences.append(s)
                        if s > 0.05:
                            texts.append(str(t))

            ocr_text = "\n".join(texts).strip() or "[无有效文字]"
            ocr_text, is_llm_failed = optimize_ocr_with_llm(ocr_text, img_file.name)
            if is_llm_failed:
                llm_failed_imgs.append(img_short_name)

            success_img_count += 1
            logger.info(f"OCR完成: {img_file.name} (识别字符数: {len(ocr_text)})")
        except Exception as e:
            logger.error(f"OCR失败 {img_file.name}: {str(e)}")
            ocr_text = "[OCR处理异常]"
            failed_imgs.append(img_short_name)

        rid_to_ocr[rId] = ocr_text
        ocr_path = get_unique_filepath(Path(OCR_DIR) / f"{img_file.stem}_ocr.txt")
        with open(ocr_path, "w", encoding="utf-8") as f:
            f.write(ocr_text)

    return rid_to_ocr, all_confidences, failed_imgs, success_img_count, llm_failed_imgs


def build_flat_doc(elements: list, ocr_map: dict, output_path: Path):
    """构建扁平化文档"""
    doc = Document()
    for elem_type, content in elements:
        if elem_type == 'text':
            doc.add_paragraph(content)
        elif elem_type == 'table':
            p = doc.add_paragraph("[表格内容]:")
            p.bold = True
            doc.add_paragraph(content)
        elif elem_type == 'image':
            p = doc.add_paragraph("[图片OCR文字]:")
            p.bold = True
            doc.add_paragraph(ocr_map.get(content, "[OCR失败]"))
    doc.save(output_path)
    logger.info(f"生成扁平化文档: {output_path.name}")
    return output_path


# 批量处理主逻辑
def process_single_doc(docx_path: Path) -> tuple[bool, str | None]:
    """处理单个文档（修复：区分纯文本文档和有图片文档的成功判定）"""
    start_time = datetime.now()
    doc_name = docx_path.name
    all_confidences = []
    failed_imgs = []
    llm_failed_imgs = []
    success_img_count = 0
    total_img_count = 0

    try:
        logger.info(f"\n{'=' * 30}")
        logger.info(f"开始处理: {doc_name}")
        prefix = get_unique_prefix(docx_path.stem)
        logger.info(f"文档前缀: {prefix}")
        elements = parse_doc_elements(docx_path)
        logger.info(f"解析完成: 共 {len(elements)} 个文档元素")
        img_map = extract_images(docx_path, prefix)
        total_img_count = len(img_map)  # 关键：记录文档实际图片总数

        # 有图片才执行OCR，无图片跳过OCR流程
        if total_img_count > 0:
            ocr_map, all_confidences, failed_imgs, success_img_count, llm_failed_imgs = ocr_images(img_map, prefix)
        else:
            ocr_map = {}  # 无图片时OCR结果为空
            logger.info("文档无嵌入式图片，无需执行OCR识别")

        elapsed_time = (datetime.now() - start_time).total_seconds()
        avg_confidence = sum(all_confidences) / len(all_confidences) if all_confidences else 1.0  # 无图片时置信度设为1.0（正常状态）
        has_valid_content = any(t in ('text', 'table') and c.strip() for t, c in elements)  # 判定是否有有效文本/表格
        has_valid_ocr = any(text not in ("[无有效文字]", "[OCR处理异常]") for text in
                            ocr_map.values()) if total_img_count > 0 else True  # 无图片时OCR有效性设为True

        # 核心修改：分情况判定成功条件
        if total_img_count == 0:
            # 纯文本文档：只要有有效文本/表格且无解析异常，即为成功
            is_success = has_valid_content
            status_desc = "纯文本文档处理成功（无图片，无需OCR）" if is_success else "纯文本文档处理失败（无有效文本/表格）"
        else:
            # 有图片文档：保持原判定逻辑
            is_success = not failed_imgs and has_valid_ocr and avg_confidence > 0.05
            status_desc = "OCR图片识别成功" if is_success else "OCR图片识别失败"

        # 构建日志内容（区分有无图片场景）
        if total_img_count == 0:
            img_desc = "文档无嵌入式图片，无需OCR识别"
        else:
            if failed_imgs:
                img_desc = f"OCR识别{success_img_count}张图片成功，{len(failed_imgs)}张图片失败（失败图片：{', '.join(failed_imgs)}）"
            else:
                img_desc = f"OCR识别{total_img_count}张图片成功"
            if llm_failed_imgs:
                img_desc += f"，大模型调用失败{len(llm_failed_imgs)}张（失败图片：{', '.join(llm_failed_imgs)}）"

        log_content = f"{doc_name} {status_desc}，耗时{elapsed_time:.1f}秒，平均置信度{avg_confidence:.4f}，{img_desc}"
        logger.info(log_content) if is_success else logger.error(log_content)

        cleaned_name = sanitize_filename(docx_path.stem)
        output_stem = f"{cleaned_name}_success" if is_success else f"{cleaned_name}_failure"
        output_path = get_unique_filepath(Path(OUTPUT_DIR) / f"{output_stem}.docx")
        build_flat_doc(elements, ocr_map, output_path)

        # 统计总字符数（包含文本、表格、OCR结果）
        text_table_chars = sum(len(c) for t, c in elements if t in ('text', 'table'))
        ocr_chars = sum(len(c) for c in ocr_map.values())
        total_chars = text_table_chars + ocr_chars
        logger.info(
            f"文档处理完成: 文本/表格字符数 {text_table_chars:,}，OCR字符数 {ocr_chars:,}，总字符数 {total_chars:,}")
        return is_success, output_path.name
    except Exception as e:
        elapsed_time = (datetime.now() - start_time).total_seconds()
        logger.error(f"处理失败 {doc_name}: {str(e)}")
        cleaned_name = sanitize_filename(docx_path.stem)
        output_stem = f"{cleaned_name}_failure"
        output_path = get_unique_filepath(Path(OUTPUT_DIR) / f"{output_stem}.docx")

        # 异常场景日志描述
        if total_img_count == 0:
            img_desc = "文档无嵌入式图片，但处理过程中发生异常"
        else:
            img_desc = f"OCR识别{success_img_count}张图片成功，{len(failed_imgs)}张图片失败（失败图片：{', '.join(failed_imgs + ['文档级处理异常'])}）"
            if llm_failed_imgs:
                img_desc += f"，大模型调用失败{len(llm_failed_imgs)}张（失败图片：{', '.join(llm_failed_imgs)}）"
        log_content = f"{doc_name} 处理异常，耗时{elapsed_time:.1f}秒，平均置信度0.0000，{img_desc}"
        logger.error(log_content)

        return False, output_path.name


def clean_process_materials():
    """清理过程材料"""
    logger.info(f"\n{'=' * 30}")
    logger.info("开始清理过程材料...")
    for dir_path in [IMAGE_DIR, OCR_DIR]:
        if Path(dir_path).exists():
            try:
                shutil.rmtree(dir_path)
                logger.info(f"已删除: {dir_path}")
            except Exception as e:
                logger.error(f"删除失败 {dir_path}: {str(e)}")


def batch_process():
    """批量处理主函数"""
    logger.info("启动DOCX文件批量处理程序")
    logger.info(f"输入目录: {INPUT_DIR}")
    logger.info(f"输出目录: {OUTPUT_DIR}")
    create_required_dirs()

    docx_files = list(Path(INPUT_DIR).glob("*.docx"))
    if not docx_files:
        logger.warning(f"未找到任何DOCX文件！请检查输入目录: {INPUT_DIR}")
        return

    logger.info(f"\n 发现 {len(docx_files)} 个待处理文件:")
    for f in docx_files:
        logger.info(f"  - {f.name}")

    success_count = 0
    results = []
    for idx, doc_path in enumerate(docx_files, 1):
        logger.info(f"\n[{idx}/{len(docx_files)}] {'-' * 20}")
        success, output_name = process_single_doc(doc_path)
        success_count += 1 if success else 0
        results.append((doc_path.name, output_name or "无", "成功" if success else "失败"))

    logger.info(f"\n{'=' * 40}")
    logger.info("批量处理最终总结:")
    logger.info(f"总文件数: {len(docx_files)} | 成功数: {success_count} | 失败数: {len(docx_files) - success_count}")
    logger.info("\n文件处理详情:")
    for orig, output, status in results:
        logger.info(f"{orig} → {output} ({status})")
    logger.info(f"\n 最终输出文件存放目录: {OUTPUT_DIR}")

    clean_process_materials()


if __name__ == "__main__":
    batch_process()
    logger.info("\n 所有处理完成！仅保留最终输出文档和错误日志。")