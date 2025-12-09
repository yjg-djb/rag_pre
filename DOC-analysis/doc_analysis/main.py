import os
import re
import shutil
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
# 修改点1：删除PaddleOCR导入
# from paddleocr import PaddleOCR
from PIL import Image, ImageEnhance
# import mammoth
# from bs4 import BeautifulSoup

# 导入其他模块
from config import INPUT_DIR, OUTPUT_DIR, IMAGE_DIR, OCR_DIR, LOG_DIR
from log import logger
from llm_client import optimize_ocr_with_llm
# 修改点2：导入第三方OCR的batch_ocr函数（二选一，根据实际场景）
from third_ocr_caseA import batch_ocr
# from third_ocr_caseB import batch_ocr

# 全局初始化
# 修改点3：删除PaddleOCR引擎初始化
# OCR_ENGINE = PaddleOCR(lang='ch')
USED_PREFIXES = set()


# 通用工具函数（无修改）
def create_required_dirs():
    """创建所有必要目录"""
    for dir_path in [INPUT_DIR, OUTPUT_DIR, IMAGE_DIR, OCR_DIR, LOG_DIR]:
        Path(dir_path).mkdir(parents=True, exist_ok=True)


def get_unique_filepath(base_path: Path) -> Path:
    """生成唯一路径避免冲突"""
    if not base_path.exists():
        return base_path
    counter = 1
    while True:
        new_path = base_path.with_stem(f"{base_path.stem}({counter})")
        if not new_path.exists():
            return new_path
        counter += 1


def sanitize_filename(name: str) -> str:
    """清理文件名非法字符"""
    return re.sub(r'[\\/*?:"<>|]', '_', name)


def get_unique_prefix(base_name: str) -> str:
    """生成唯一前缀"""
    prefix = base_name[:50]
    unique_prefix = prefix
    counter = 1
    while unique_prefix in USED_PREFIXES:
        unique_prefix = f"{prefix}_{counter}"
        counter += 1
    USED_PREFIXES.add(unique_prefix)
    return unique_prefix


# 核心处理函数（表格相关函数无修改）
def parse_cell_text(cell_elem):
    """解析单元格内的文本内容，处理多行文本"""
    texts = []
    for para in cell_elem.iterfind('.//w:p', namespaces=cell_elem.nsmap):
        para_text = ''.join(
            run.text for run in para.iterfind('.//w:t', namespaces=cell_elem.nsmap) if run.text
        ).strip()
        if para_text:
            texts.append(para_text)
    return '\n'.join(texts)


def is_header_row(row_elem):
    """判断行是否为表头行"""
    # 检查行内是否有表头标记或样式
    for cell in row_elem.iterfind('.//w:tc', namespaces=row_elem.nsmap):
        for p in cell.iterfind('.//w:p', namespaces=cell.nsmap):  # 修复原代码笔误：cell.elem → cell
            p_pr = p.find('.//w:pPr', namespaces=p.nsmap)
            if p_pr is not None:
                p_style = p_pr.find('.//w:pStyle', namespaces=p_pr.nsmap)
                if p_style is not None and p_style.get(qn('w:val')) == '表头':
                    return True
    # 检查表格是否有表头定义
    tbl = row_elem.getparent()
    tbl_pr = tbl.find('.//w:tblPr', namespaces=tbl.nsmap)
    if tbl_pr is not None:
        tbl_header = tbl_pr.find('.//w:tblHeader', namespaces=tbl_pr.nsmap)
        if tbl_header is not None:
            return True
    return False


def process_merged_cells(table_elem):
    """处理合并单元格，构建完整表格数据"""
    nsmap = table_elem.nsmap
    rows = list(table_elem.iterfind('.//w:tr', namespaces=nsmap))
    if not rows:
        return []

    # 获取表格列数
    grid_cols = table_elem.find('.//w:tblGrid/w:gridCol', namespaces=nsmap)
    if grid_cols is None:
        cols_count = max(len(list(row.iterfind('.//w:tc', namespaces=nsmap))) for row in rows)
    else:
        cols_count = len(table_elem.findall('.//w:tblGrid/w:gridCol', namespaces=nsmap))

    # 初始化表格矩阵
    table_data = []
    merged_cells = []  # 记录合并单元格信息 (row, col, width, height, value)

    for row_idx, row in enumerate(rows):
        row_data = []
        cells = list(row.iterfind('.//w:tc', namespaces=nsmap))
        col_idx = 0

        # 跳过已被合并的单元格
        while any(m[0] == row_idx and m[1] == col_idx for m in merged_cells):
            col_idx += 1
            if col_idx >= cols_count:
                break

        for cell in cells:
            if col_idx >= cols_count:
                break

            # 检查是否为合并单元格
            grid_span = cell.find('.//w:tcPr/w:gridSpan', namespaces=nsmap)
            v_merge = cell.find('.//w:tcPr/w:vMerge', namespaces=nsmap)

            col_span = int(grid_span.get(qn('w:val'))) if grid_span is not None else 1
            row_span = 1

            # 处理行合并
            if v_merge is not None:
                v_merge_val = v_merge.get(qn('w:val'))
                if v_merge_val == 'restart':
                    # 查找合并的行数
                    current_row = row_idx + 1
                    while current_row < len(rows):
                        next_cell = None
                        # 查找下一行中对应位置的单元格
                        next_row_cells = list(rows[current_row].iterfind('.//w:tc', namespaces=nsmap))
                        temp_col = col_idx
                        cell_count = 0
                        for c in next_row_cells:
                            c_span = int(c.find('.//w:tcPr/w:gridSpan', namespaces=nsmap).get(qn('w:val'))) if c.find(
                                './/w:tcPr/w:gridSpan', namespaces=nsmap) is not None else 1
                            if temp_col < cell_count + c_span:
                                next_cell = c
                                break
                            cell_count += c_span

                        if next_cell is not None and next_cell.find('.//w:tcPr/w:vMerge', namespaces=nsmap) is not None:
                            row_span += 1
                            current_row += 1
                        else:
                            break
                    merged_cells.append((row_idx, col_idx, col_span, row_span, parse_cell_text(cell)))

            # 填充单元格内容
            cell_text = parse_cell_text(cell)
            row_data.append(cell_text)

            # 处理列合并，填充空内容
            for i in range(1, col_span):
                if col_idx + i < cols_count:
                    row_data.append('')

            # 更新列索引
            col_idx += col_span

            # 跳过已被合并的单元格
            while any(m[0] == row_idx and m[1] == col_idx for m in merged_cells):
                row_data.append('')
                col_idx += 1
                if col_idx >= cols_count:
                    break

        # 补充行尾可能缺失的单元格
        while len(row_data) < cols_count:
            row_data.append('')

        table_data.append(row_data)

    # 填充合并单元格的内容
    for (start_row, start_col, width, height, value) in merged_cells:
        for i in range(height):
            for j in range(width):
                if start_row + i < len(table_data) and start_col + j < len(table_data[start_row + i]):
                    table_data[start_row + i][start_col + j] = value

    return table_data


def extract_table_text(docx_path: Path) -> list:
    """优化后的表格提取函数，通过解析XML处理合并单元格和表头"""
    try:
        doc = Document(docx_path)
        table_texts = []

        for table in doc.tables:
            # 获取底层XML元素
            table_elem = table._element
            table_data = process_merged_cells(table_elem)

            if not table_data:
                continue

            # 识别表头行
            header_rows = []
            data_rows = []
            for i, row in enumerate(table_data):
                if is_header_row(table_elem[i]):  # table_elem[i] 是第i行的XML元素
                    header_rows.append(row)
                else:
                    data_rows.append(row)

            # 处理无表头的情况
            if not header_rows:
                header_rows = [table_data[0]]
                data_rows = table_data[1:] if len(table_data) > 1 else []

            # 格式化输出
            formatted_table = []
            if header_rows:
                # 合并可能存在的多行表头
                headers = []
                for header_row in header_rows:
                    for i, header in enumerate(header_row):
                        if i < len(headers):
                            headers[i] = f"{headers[i]} {header}".strip()
                        else:
                            headers.append(header.strip())

                # 处理数据行
                for row in data_rows:
                    row_str = []
                    for i, (header, value) in enumerate(zip(headers, row)):
                        if header:  # 只有表头存在时才添加键值对格式
                            row_str.append(f"{header}: {value}" if value else f"{header}: ")
                        else:
                            row_str.append(value)
                    if row_str:
                        formatted_table.append(", ".join(row_str))
            else:
                # 无表头表格直接拼接
                for row in table_data:
                    if any(cell.strip() for cell in row):
                        formatted_table.append(", ".join(row))

            if formatted_table:
                table_texts.append("\n".join(formatted_table))

        return table_texts
    except Exception as e:
        logger.error(f"表格提取失败 {docx_path.name}: {str(e)}")
        return []

def parse_doc_elements(docx_path: Path) -> list:
    """解析文档元素（无修改）"""
    try:
        table_texts = extract_table_text(docx_path)
        doc = Document(docx_path)
        elements = []
        table_idx = 0

        for elem in doc._element.iter():
            if elem.tag.endswith('}p') and not any(parent.tag.endswith('}tbl') for parent in elem.iterancestors()):
                text = ''.join(run.text for run in elem.iter(qn('w:r')) if run.text).strip()
                if text:
                    elements.append(('text', text))
                for blip in elem.iter(qn('a:blip')):
                    if rId := blip.get(qn('r:embed')):
                        elements.append(('image', rId))
            elif elem.tag.endswith('}tbl') and table_idx < len(table_texts):
                elements.append(('table', table_texts[table_idx]))
                table_idx += 1
        return elements
    except Exception as e:
        logger.error(f"文档解析失败 {docx_path.name}: {str(e)}")
        raise


def extract_images(docx_path: Path, prefix: str) -> dict:
    """提取图片（无修改）"""
    doc = Document(docx_path)
    image_rels = {rel.rId: rel.target_part.blob for rel in doc.part.rels.values() if "image" in rel.target_ref}
    rid_to_img = {}

    for i, (rId, blob) in enumerate(image_rels.items(), 1):
        base_path = Path(IMAGE_DIR) / f"{prefix}_img_{i:03d}.png"
        img_path = get_unique_filepath(base_path)
        with open(img_path, "wb") as f:
            f.write(blob)
        rid_to_img[rId] = str(img_path)

    logger.info(f"提取图片 {len(rid_to_img)} 张")
    return rid_to_img


def enhance_image(img_path: str) -> str:
    """增强图片对比度（无修改）"""
    try:
        img = Image.open(img_path)
        enhanced = ImageEnhance.Contrast(img).enhance(2.0)
        enhanced_path = get_unique_filepath(Path(img_path).with_stem(f"{Path(img_path).stem}_enhanced"))
        enhanced.save(enhanced_path)
        logger.info(f"增强图片对比度: {Path(img_path).name}")
        return str(enhanced_path)
    except Exception as e:
        logger.error(f"图片增强失败: {e}")
        return img_path


"""修改点：删除所有置信度相关逻辑"""
def ocr_images(image_map: dict, prefix: str) -> tuple[dict, list[str], int, list[str]]:
    """修改点：删除置信度列表all_confidences"""
    rid_to_ocr = {}
    failed_imgs = []
    llm_failed_imgs = []
    success_img_count = 0

    # 步骤1：构造第三方OCR需要的(图片路径, 图片名称)元组列表
    img_info_list = []
    rId_list = []  # 记录rId顺序，匹配第三方返回结果
    for rId, img_path in image_map.items():
        img_file = Path(img_path)
        img_name = img_file.name  # 图片名称（传给第三方接口）
        img_info_list.append((img_path, img_name))
        rId_list.append(rId)

    # 步骤2：调用第三方OCR批量识别
    try:
        print("第三方OCR返回结果")
        third_ocr_results = batch_ocr(img_info_list)
        print(f"第三方OCR返回结果类型：{type(third_ocr_results)}")
        # logger.info(f"第三方OCR返回结果类型：{type(third_ocr_results)}")
        # logger.info(f"第三方OCR返回结果：{third_ocr_results}")
    except Exception as e:
        logger.error(f"调用第三方OCR接口批量失败: {str(e)}")
        for rId, img_path in image_map.items():
            img_file = Path(img_path)
            img_short_name = img_file.name.split("_")[-1]
            rid_to_ocr[rId] = "[OCR处理异常]"
            failed_imgs.append(img_short_name)
        return rid_to_ocr, failed_imgs, success_img_count, llm_failed_imgs  # 修改点：删除all_confidences

    # 步骤3：解析第三方OCR结果
    for idx, (rId, third_result) in enumerate(zip(rId_list, third_ocr_results)):
        img_path = img_info_list[idx][0]
        img_file = Path(img_path)
        img_short_name = img_file.name.split("_")[-1]
        logger.info(f"解析第三方OCR结果: {img_file.name}")

        try:
            # 提取文本（删除置信度相关代码）
            texts = []
            for text_block in third_result:
                if len(text_block) >= 2 and len(text_block[1]) >= 1:
                    t = text_block[1][0].strip()
                    if t:
                        texts.append(t)

            ocr_text = "\n".join(texts).strip() or "[无有效文字]"
            # 调用LLM优化
            logger.info(f"原本txt{ocr_text}")
            ocr_text, is_llm_failed = optimize_ocr_with_llm(ocr_text, img_file.name)
            if is_llm_failed:
                llm_failed_imgs.append(img_short_name)

            success_img_count += 1  # 删除置信度添加逻辑
            logger.info(f"第三方OCR完成: {img_file.name} (识别字符数: {len(ocr_text)})")
        except Exception as e:
            logger.error(f"解析第三方OCR结果失败 {img_file.name}: {str(e)}")
            ocr_text = "[OCR处理异常]"
            failed_imgs.append(img_short_name)

        rid_to_ocr[rId] = ocr_text
        ocr_path = get_unique_filepath(Path(OCR_DIR) / f"{img_file.stem}_ocr.txt")
        with open(ocr_path, "w", encoding="utf-8") as f:
            f.write(ocr_text)

    return rid_to_ocr, failed_imgs, success_img_count, llm_failed_imgs  # 修改点：删除all_confidences

def build_flat_doc(elements: list, ocr_map: dict, output_path: Path):
    """构建扁平化文档（无修改）"""
    doc = Document()
    for elem_type, content in elements:
        if elem_type == 'text':
            doc.add_paragraph(content)
        elif elem_type == 'table':
            p = doc.add_paragraph("[表格内容]:")
            p.bold = True
            doc.add_paragraph(content)
        elif elem_type == 'image':
            p = doc.add_paragraph("[图片OCR文字]:")
            p.bold = True
            doc.add_paragraph(ocr_map.get(content, "[OCR失败]"))
    doc.save(output_path)
    logger.info(f"生成扁平化文档: {output_path.name}")
    return output_path


# 批量处理主逻辑（无修改，仅兼容置信度逻辑）
"""处理单个文档（删除所有置信度相关内容）"""
def process_single_doc(docx_path: Path) -> tuple[bool, str | None]:
    start_time = datetime.now()
    doc_name = docx_path.name
    failed_imgs = []
    llm_failed_imgs = []
    success_img_count = 0
    total_img_count = 0

    try:
        logger.info(f"\n{'=' * 30}")
        logger.info(f"开始处理: {doc_name}")
        prefix = get_unique_prefix(docx_path.stem)
        logger.info(f"文档前缀: {prefix}")
        elements = parse_doc_elements(docx_path)
        logger.info(f"解析完成: 共 {len(elements)} 个文档元素")
        img_map = extract_images(docx_path, prefix)
        total_img_count = len(img_map)

        # 有图片才执行OCR
        if total_img_count > 0:
            # 修改点：接收参数删除all_confidences
            ocr_map, failed_imgs, success_img_count, llm_failed_imgs = ocr_images(img_map, prefix)
        else:
            ocr_map = {}
            logger.info("文档无嵌入式图片，无需执行OCR识别")

        elapsed_time = (datetime.now() - start_time).total_seconds()
        has_valid_content = any(t in ('text', 'table') and c.strip() for t, c in elements)

        has_valid_ocr = any(text not in ("[无有效文字]", "[OCR处理异常]") for text in
                            ocr_map.values()) if total_img_count > 0 else True

        # 核心修改：删除置信度相关的成功条件判断
        if total_img_count == 0:
            is_success = has_valid_content
            status_desc = "纯文本文档处理成功（无图片，无需OCR）" if is_success else "纯文本文档处理失败（无有效文本/表格）"
        else:
            # 仅保留图片处理状态和有效OCR的判断
            is_success = not failed_imgs and has_valid_ocr
            status_desc = "OCR图片识别成功" if is_success else "OCR图片识别失败"

        # 构建日志内容（删除平均置信度）
        if total_img_count == 0:
            img_desc = "文档无嵌入式图片，无需OCR识别"
        else:
            if failed_imgs:
                img_desc = f"OCR识别{success_img_count}张图片成功，{len(failed_imgs)}张图片失败（失败图片：{', '.join(failed_imgs)}）"
            else:
                img_desc = f"OCR识别{total_img_count}张图片成功"
            if llm_failed_imgs:
                img_desc += f"，大模型调用失败{len(llm_failed_imgs)}张（失败图片：{', '.join(llm_failed_imgs)}）"

        # 修改点：日志删除平均置信度
        log_content = f"{doc_name} {status_desc}，耗时{elapsed_time:.1f}秒，{img_desc}"
        logger.info(log_content) if is_success else logger.error(log_content)

        cleaned_name = sanitize_filename(docx_path.stem)
        output_stem = f"{cleaned_name}_success" if is_success else f"{cleaned_name}_failure"
        output_path = get_unique_filepath(Path(OUTPUT_DIR) / f"{output_stem}.docx")
        build_flat_doc(elements, ocr_map, output_path)

        # 统计字符数（保留原有逻辑）
        text_table_chars = sum(len(c) for t, c in elements if t in ('text', 'table'))
        ocr_chars = sum(len(c) for c in ocr_map.values())
        total_chars = text_table_chars + ocr_chars
        logger.info(
            f"文档处理完成: 文本/表格字符数 {text_table_chars:,}，OCR字符数 {ocr_chars:,}，总字符数 {total_chars:,}")
        return is_success, output_path.name
    except Exception as e:
        elapsed_time = (datetime.now() - start_time).total_seconds()
        logger.error(f"处理失败 {doc_name}: {str(e)}")
        cleaned_name = sanitize_filename(docx_path.stem)
        output_stem = f"{cleaned_name}_failure"
        output_path = get_unique_filepath(Path(OUTPUT_DIR) / f"{output_stem}.docx")

        # 异常日志删除平均置信度
        if total_img_count == 0:
            img_desc = "文档无嵌入式图片，但处理过程中发生异常"
        else:
            img_desc = f"OCR识别{success_img_count}张图片成功，{len(failed_imgs)}张图片失败（失败图片：{', '.join(failed_imgs + ['文档级处理异常'])}）"
            if llm_failed_imgs:
                img_desc += f"，大模型调用失败{len(llm_failed_imgs)}张（失败图片：{', '.join(llm_failed_imgs)}）"
        log_content = f"{doc_name} 处理异常，耗时{elapsed_time:.1f}秒，{img_desc}"  # 删除平均置信度
        logger.error(log_content)

        return False, output_path.name

def clean_process_materials():
    """清理过程材料（无修改）"""
    logger.info(f"\n{'=' * 30}")
    logger.info("开始清理过程材料...")
    for dir_path in [IMAGE_DIR, OCR_DIR]:
        if Path(dir_path).exists():
            try:
                shutil.rmtree(dir_path)
                logger.info(f"已删除: {dir_path}")
            except Exception as e:
                logger.error(f"删除失败 {dir_path}: {str(e)}")


def batch_process():
    """批量处理主函数（无修改）"""
    logger.info("启动DOCX文件批量处理程序")
    logger.info(f"输入目录: {INPUT_DIR}")
    logger.info(f"输出目录: {OUTPUT_DIR}")
    create_required_dirs()

    docx_files = list(Path(INPUT_DIR).glob("*.docx"))
    if not docx_files:
        logger.warning(f"未找到任何DOCX文件！请检查输入目录: {INPUT_DIR}")
        return

    logger.info(f"\n 发现 {len(docx_files)} 个待处理文件:")
    for f in docx_files:
        logger.info(f"  - {f.name}")

    success_count = 0
    results = []
    for idx, doc_path in enumerate(docx_files, 1):
        logger.info(f"\n[{idx}/{len(docx_files)}] {'-' * 20}")
        success, output_name = process_single_doc(doc_path)
        success_count += 1 if success else 0
        results.append((doc_path.name, output_name or "无", "成功" if success else "失败"))

    logger.info(f"\n{'=' * 40}")
    logger.info("批量处理最终总结:")
    logger.info(f"总文件数: {len(docx_files)} | 成功数: {success_count} | 失败数: {len(docx_files) - success_count}")
    logger.info("\n文件处理详情:")
    for orig, output, status in results:
        logger.info(f"{orig} → {output} ({status})")
    logger.info(f"\n 最终输出文件存放目录: {OUTPUT_DIR}")

    # clean_process_materials()


if __name__ == "__main__":
    batch_process()
    logger.info("\n 所有处理完成！仅保留最终输出文档和错误日志。")